from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT.parent / "output" / "sop"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "LMS_Move_Request_SOP.docx"

NAVY = RGBColor(31, 47, 66)
TEAL = RGBColor(73, 102, 115)
SLATE = RGBColor(92, 106, 125)
LIGHT = "EAF0F3"
PALE = "F6F8F9"
GREEN = "E7F3EC"
AMBER = "FFF4D8"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep(paragraph, next_=False):
    ppr = paragraph._p.get_or_add_pPr()
    tag = "keepNext" if next_ else "keepLines"
    if ppr.find(qn(f"w:{tag}")) is None:
        ppr.append(OxmlElement(f"w:{tag}"))


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    keep(p, True)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_note(doc, title, body, fill=AMBER):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell, 120, 150, 120, 150)
    p = cell.paragraphs[0]
    r = p.add_run(title + " ")
    r.bold = True
    r.font.color.rgb = NAVY
    p.add_run(body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_snapshot(doc, filename, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.add_run().add_picture(str(ROOT / filename), width=Inches(4.25))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.style = doc.styles["Caption"]
    c.add_run("  Training example—sample data only; not a live transaction.")


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.78)
sec.right_margin = Inches(0.78)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.3)
normal.font.color.rgb = NAVY
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for name, size, color, before, after in (
    ("Title", 28, NAVY, 0, 10),
    ("Subtitle", 15, TEAL, 0, 14),
    ("Heading 1", 18, NAVY, 14, 7),
    ("Heading 2", 13.5, TEAL, 11, 5),
    ("Heading 3", 11, SLATE, 8, 3),
):
    s = styles[name]
    s.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
    s.font.size = Pt(size)
    s.font.bold = name != "Subtitle"
    s.font.color.rgb = color
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)

styles["Caption"].font.name = "Aptos"
styles["Caption"].font.size = Pt(8.5)
styles["Caption"].font.italic = True
styles["Caption"].font.color.rgb = SLATE

header = sec.header.paragraphs[0]
header.text = "LMS IMPORTS  |  STANDARD OPERATING PROCEDURE"
header.style = styles["Caption"]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer = sec.footer.paragraphs[0]
footer.add_run("LMS-TRK-001  •  Internal Training Guide                         ")
page_number(footer)

# Editorial cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(70)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LMS IMPORTS")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = TEAL
p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Move Request\nStandard Operating Procedure")
p = doc.add_paragraph(style="Subtitle")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Mobile and desktop requestor workflow—from new request through dispatch")

t = doc.add_table(rows=5, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
for row in t.rows:
    row.cells[0].width = Inches(1.7)
    row.cells[1].width = Inches(4.2)
metadata = [
    ("SOP ID", "LMS-TRK-001"),
    ("Effective date", "August 6, 2026"),
    ("Process owner", "LMS Trucking / Dispatch"),
    ("Audience", "Requestors, dispatchers, trucking drivers, and administrators"),
    ("System", "LMS Imports — Trucking Move Requests"),
]
for i, (a, b) in enumerate(metadata):
    shade(t.cell(i, 0), LIGHT)
    for c in t.rows[i].cells: margins(c, 100, 120, 100, 120)
    t.cell(i, 0).paragraphs[0].add_run(a).bold = True
    t.cell(i, 1).paragraphs[0].add_run(b)

doc.add_paragraph()
add_note(doc, "Document use:", "Follow this SOP when submitting a new trucking move. Screenshots use sample information and must not be copied into a live request.", GREEN)
doc.add_page_break()

add_heading(doc, "1. Purpose", 1)
doc.add_paragraph("This SOP explains how to create a complete, dispatch-ready Move Request in LMS Imports. It standardizes required information, service-specific fields, validation, submission, status tracking, editing, and cancellation so dispatch receives enough information to schedule the work correctly.")

add_heading(doc, "2. Scope", 1)
doc.add_paragraph("Use this procedure for trucking move requests entered from the desktop application or the requestor mobile view. It covers request creation through submission and tracking. Dispatcher assignment, driver execution, signatures, tickets, and billing occur after the request is submitted and are summarized in Section 10.")
add_note(doc, "Important:", "The trucking workflow supports operational monitoring and billing preparation. It does not by itself authorize users to post accounting entries in another accounting system.")

add_heading(doc, "3. Roles and responsibilities", 1)
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
for i, txt in enumerate(("Role", "Responsibility")):
    shade(table.cell(0, i), LIGHT); table.cell(0, i).paragraphs[0].add_run(txt).bold = True
roles = [
    ("Requestor", "Selects the correct customer, jobsite, service, timing, payment mode, email recipient, and move details; reviews all information before saving."),
    ("Dispatcher", "Reviews unscheduled requests; assigns the dispatch date, trucking driver, and actual truck/equipment; may cancel with a documented reason."),
    ("Driver", "Reviews assigned work, records permitted final details, captures required driver and customer signatures, and finalizes the task."),
    ("Administrator", "Maintains service patterns, rate-sheet options, users, customer/jobsite master data, and equipment master records."),
]
for role, resp in roles:
    cells = table.add_row().cells
    cells[0].text = role; cells[1].text = resp
    for c in cells: margins(c)

add_heading(doc, "4. Before you begin", 1)
for text in [
    "Confirm you are using the correct login and have access to Equipment Request and/or Trucking Move Request as assigned.",
    "Collect the customer, project/jobsite, requested date, AM/PM preference, service type, payment method, and a working email address.",
    "For equipment or container moves, have the equipment identifier, plate/VIN/asset number, container/trailer number, size, and routing information available.",
    "If a master-data option is missing, contact an administrator. Do not select an inaccurate substitute merely to submit the request.",
]: add_bullet(doc, text)

doc.add_page_break()
add_heading(doc, "5. Create a new Move Request", 1)

add_heading(doc, "5.1 Open the form", 2)
add_number(doc, "Open the Trucking module and select Move Requests. In the mobile requestor view, select Trucking Move Request.")
add_number(doc, "Select New request. The New Move Request window opens.")
add_number(doc, "If a tutorial button is available, select it for field-by-field guidance. Close or advance the tutorial when ready to enter data.")

add_heading(doc, "5.2 Confirm request identification", 2)
add_bullet(doc, "Request # is generated by the system in the LMS-00001 sequence. Do not type over it.")
add_bullet(doc, "Entry Date is fixed to the date the request is created. It is not the requested delivery date.")
add_bullet(doc, "A new number should be used for each separate request. Do not reuse a prior request number.")

add_heading(doc, "5.3 Enter customer and contact information", 2)
steps = [
    "Customer — select the customer from Customer Master. Verify the legal/customer name before continuing.",
    "Project / Jobsite — type to search the available Fleet & Equipment jobsites/locations, then select the correct result. If permitted and the jobsite is missing, enter a new value; otherwise request administrator assistance.",
    "Contact # — review the number populated from Customer Master, if available, and correct it when the request requires a different contact.",
    "Requested Delivery Date — choose the date the customer wants the service. This is separate from the fixed Entry Date.",
    "Email — enter the requestor/customer email that must receive notifications. Email is required. When a shared account such as request@lmsfm.com is used to log in, notifications still go to the email entered on the request.",
]
for s in steps: add_number(doc, s)
add_snapshot(doc, "move-request-1-full.png", "Figure 1. Request identification, customer, jobsite, contact, delivery date, and required email.")

add_heading(doc, "5.4 Enter delivery and routing details", 2)
for s in [
    "Delivery Time — select AM or PM.",
    "Trip Type — select One-Way (1×) or Round Trip (2×). Round Trip doubles fixed-rate charges only; it does not double an hourly rate.",
    "Origin — search and select a jobsite/location when known. Origin is optional unless the selected service or operating procedure requires it.",
    "Destination — search and select a jobsite/location when known. Destination is optional unless the selected service or operating procedure requires it.",
]: add_number(doc, s)

add_heading(doc, "5.5 Select Service Needed", 2)
doc.add_paragraph("Select the top-level service that describes the work. The form must reveal only the fields needed for that service. Complete every service-specific field that appears; hidden fields are not required for the selected service.")
add_snapshot(doc, "move-request-2-full.png", "Figure 2. Delivery time, trip type, service, and conditional service fields.")

doc.add_page_break()
add_heading(doc, "6. Service-specific requirements", 1)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, txt in enumerate(("Service", "Fields that appear", "How to complete")):
    shade(table.cell(0, i), LIGHT); table.cell(0, i).paragraphs[0].add_run(txt).bold = True
services = [
    ("Equipment Move", "Equipment lookup", "Search Equipment Master by asset number, description, VIN, or plate number and select the exact equipment. Only an administrator should create a missing equipment record."),
    ("Container / Trailer Move", "Move type; size; container/trailer #", "Choose Sidelifter or Chassis Move; select 20, 40, 45, or Double Wide; enter the container/trailer reference."),
    ("High Deck Trailer; Highbed/Chameleon; Flatrack/Chameleon", "Material(s); unit of measurement; weight (lb); dimensions", "Describe the load accurately, select/enter the unit, and record total weight and dimensions for planning and safety."),
    ("Dump Truck; End Dump", "Materials loaded; # of loads; CY/load; total CY", "Enter loads and cubic yards per load. Verify Total CY = # of Loads × CY/Load."),
    ("Water Service", "Configured water capacity", "Select the configured capacity group, such as 0–1,000 gallons or 2,500–4,000 gallons, matching the rate sheet."),
    ("Roll Off Service", "Bin size/capacity; tons when known; tipping fee", "Tipping fee applies only to Roll Off. When tons are recorded, tipping is tons × rate; otherwise the configured bin-size basis applies."),
    ("Pump Service", "Configured pump-service details", "Complete the fields shown for the chosen pump service and explain the work in Move Description."),
    ("Waste Water Pump Services", "Portable Toilets; Portable Sink; tank capacity", "Select the requested unit or a configured tank capacity from 300 through 1,000 gallons in 50-gallon increments."),
]
for a, b, c in services:
    row = table.add_row().cells
    row[0].text = a; row[1].text = b; row[2].text = c
    for cell in row: margins(cell, 70, 75, 70, 75); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

add_snapshot(doc, "move-request-3-full.png", "Figure 3. Examples of service-dependent information.")
add_note(doc, "Stop and correct:", "If the form shows fields for an unrelated service, or the correct option is missing, do not improvise. Save only after the selected service and visible fields accurately describe the move.")

add_heading(doc, "7. Payment, description, and submission", 1)
add_heading(doc, "7.1 Select mode of payment", 2)
add_bullet(doc, "Cash — select when the customer will pay by cash.")
add_bullet(doc, "Credit Card — select when payment will be by credit card.")
add_bullet(doc, "PO — select when the customer is using a purchase order. PO # becomes mandatory; enter the complete customer PO number.")

add_heading(doc, "7.2 Complete the Move Description", 2)
doc.add_paragraph("Move Description is required and is vital operational information. It follows the request into scheduling and the ticket. Write a concise but complete description that answers what is moving, where it is going, timing expectations, quantities, and any special instructions.")
add_note(doc, "Good example:", "“Move one 40 ft container from Andersen Housing to Port Yard. Customer requests AM delivery. Contact site supervisor before arrival.”", GREEN)
add_note(doc, "Avoid:", "Descriptions such as “move,” “truck needed,” or “see customer.” They do not provide enough information for dispatch or the ticket.")

add_heading(doc, "7.3 Final review and save", 2)
checklist = [
    "Generated Request # and fixed Entry Date are visible.",
    "Customer and Project/Jobsite are correct.",
    "Requested Delivery Date and AM/PM are selected.",
    "Required email is correct and belongs to the intended recipient.",
    "Trip Type and Service Needed are correct.",
    "All visible service-dependent fields are complete.",
    "Payment mode is selected; PO # is present when payment mode is PO.",
    "Move Description is complete and specific.",
]
for item in checklist: add_bullet(doc, "☐ " + item)
add_number(doc, "Select Save / Submit Request once.")
add_number(doc, "Wait for the submission-success confirmation. Do not repeatedly press Save while the request is processing.")
add_number(doc, "Close the confirmation or return to the Move Requests list. Confirm the request appears under Unscheduled.")
add_snapshot(doc, "move-request-4-full.png", "Figure 4. Payment, required description, submission, and initial status.")

doc.add_page_break()
add_heading(doc, "8. Request status and tracking", 1)
statuses = [
    ("Unscheduled", "Request saved but not yet assigned by dispatch.", "Requestor/dispatcher may review or edit authorized details."),
    ("Scheduled", "Dispatch date, trucking driver, and actual equipment have been assigned.", "Dispatcher manages assignment; driver receives the task."),
    ("Assigned Driver", "Task is visible to the assigned trucking driver.", "Driver reviews details, records permitted final data, and saves progress."),
    ("Finalized / Ticket", "Work is accepted with required driver and customer signatures.", "The completed move appears in Tickets for operational record and billing preparation."),
    ("Cancelled", "Move will not proceed.", "Cancellation reason is required and the record remains traceable."),
]
table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, txt in enumerate(("Status", "Meaning", "Expected action")):
    shade(table.cell(0, i), LIGHT); table.cell(0, i).paragraphs[0].add_run(txt).bold = True
for a,b,c in statuses:
    cells=table.add_row().cells; cells[0].text=a; cells[1].text=b; cells[2].text=c
    for cell in cells: margins(cell)

add_heading(doc, "9. Edit or cancel a request", 1)
add_heading(doc, "9.1 Edit", 2)
add_number(doc, "Locate the request in the correct status tab and select Edit.")
add_number(doc, "Confirm the existing request number, signature/reference data, equipment lines, and entered information load before changing anything.")
add_number(doc, "Change only the necessary fields, recheck the required email and Move Description, then select Save.")
add_number(doc, "If existing lines or signatures disappear before you save, select Cancel/Close and report the problem; do not overwrite a complete request with a partially loaded form.")

add_heading(doc, "9.2 Cancel", 2)
add_number(doc, "Select Cancelled/Cancel Request from the request actions.")
add_number(doc, "Enter a clear cancellation reason, such as customer cancellation, duplicate request, weather, or equipment unavailable.")
add_number(doc, "Confirm the action and verify the record moves to the Cancelled tab.")
add_note(doc, "Do not delete:", "Cancellation preserves the request history and reason. Use Cancelled instead of deleting a submitted transaction.")

add_heading(doc, "10. What happens after submission", 1)
for text in [
    "Dispatch opens the unscheduled request and assigns the dispatch date, a user configured as a Trucking Driver, and the actual truck/equipment from Equipment Master.",
    "The assigned driver sees the task in Assigned Driver, can save permitted progress and final operational details, and records the final ticket date where authorized.",
    "Both driver and customer signatures are required before finalization.",
    "Finalized work becomes a ticket. The ticket carries the jobsite, Move Description, service details, dates/times, and signatures.",
    "Billing amounts and tipping-fee calculations belong on the billing summary, not on the customer-facing operational ticket where the configured ticket format suppresses them.",
]: add_bullet(doc, text)

add_heading(doc, "11. Troubleshooting", 1)
issues = [
    ("New request does not open", "Refresh once, confirm module permission, then reopen Move Requests. Report persistent blank/non-opening forms with the time and user email."),
    ("Customer or jobsite is missing", "Search using a shorter name. Confirm the record exists in Customer Master or Fleet & Equipment; request administrator creation if permitted."),
    ("Equipment or plate cannot be found", "Search by asset number, plate, VIN, and a partial description. Confirm the latest Fleet & Equipment record has synchronized before requesting admin assistance."),
    ("Save does nothing", "Review all required fields, especially Email, Move Description, delivery date, service fields, payment mode, and PO # when PO is selected."),
    ("Wrong fields are displayed", "Re-select Service Needed. If unrelated fields remain, close without saving and report the service and screenshot."),
    ("Duplicate request appears", "Do not submit again. Notify dispatch/admin and cancel the duplicate with a reason."),
]
table = doc.add_table(rows=1, cols=2); table.style="Table Grid"; table.alignment=WD_TABLE_ALIGNMENT.CENTER
for i, txt in enumerate(("Problem", "Action")):
    shade(table.cell(0,i), LIGHT); table.cell(0,i).paragraphs[0].add_run(txt).bold=True
for a,b in issues:
    row=table.add_row().cells; row[0].text=a; row[1].text=b
    for c in row: margins(c)

add_heading(doc, "12. Quick-reference submission checklist", 1)
for item in checklist: add_bullet(doc, "☐ " + item)
add_bullet(doc, "☐ Submission-success message appeared.")
add_bullet(doc, "☐ Request is visible under Unscheduled with the expected LMS request number.")

doc.add_paragraph()
add_note(doc, "End of SOP:", "For training questions or system issues, provide the request number, user email, date/time, selected service, and a screenshot to the LMS system administrator.", GREEN)

# Repeat header/footer in all sections and set document metadata.
doc.core_properties.title = "LMS Move Request Standard Operating Procedure"
doc.core_properties.subject = "Detailed SOP for creating and tracking trucking move requests"
doc.core_properties.author = "LMS Imports"
doc.core_properties.keywords = "LMS, trucking, move request, SOP, training"

doc.save(DOCX)
print(DOCX)
