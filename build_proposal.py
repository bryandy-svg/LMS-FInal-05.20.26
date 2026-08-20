from pathlib import Path
from datetime import date
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "DZSP21_Equipment_Substitution_Proposal.docx"
ASSET = ROOT / "proposal_assets"
ASSET.mkdir(exist_ok=True)

NAVY = "17365D"
BLUE = "2E74B5"
PALE = "E8EEF5"
LIGHT = "F4F6F9"
GRAY = "5C6773"
INK = "1B2430"
GOLD = "C8962E"
WHITE = "FFFFFF"
GREEN = "2D6A4F"
RED = "9B1C1C"

def font(size=11, bold=False):
    try:
        return ImageFont.truetype("arialbd.ttf" if bold else "arial.ttf", size)
    except Exception:
        return ImageFont.load_default()

def draw_forklift(path, title, electric=False, side=False):
    im = Image.new("RGB", (1500, 650), "white")
    d = ImageDraw.Draw(im)
    line = (23, 54, 93)
    fill = (232, 238, 245)
    # ground
    d.line((80, 520, 1420, 520), fill=(165,175,185), width=3)
    # chassis and counterweight
    d.rounded_rectangle((470, 330, 1050, 490), radius=35, outline=line, width=7, fill=fill)
    d.polygon([(910,330),(1110,370),(1125,465),(1020,490),(930,455)], outline=line, fill=(215,224,235))
    # overhead guard/cab
    d.line((620,330,650,155), fill=line, width=8)
    d.line((650,155,880,155), fill=line, width=8)
    d.line((880,155,930,330), fill=line, width=8)
    d.line((670,210,890,210), fill=line, width=5)
    # mast and forks
    d.rectangle((385,125,425,500), outline=line, width=6, fill=(245,247,250))
    d.rectangle((435,150,458,500), outline=line, width=5, fill=(245,247,250))
    d.line((395,445,175,445), fill=line, width=10)
    d.line((175,445,105,480), fill=line, width=8)
    # wheels
    for x, r in [(555,75),(965,65)]:
        d.ellipse((x-r,455-r,x+r,455+r), fill=(40,45,50), outline=line, width=4)
        d.ellipse((x-r//2,455-r//2,x+r//2,455+r//2), fill="white", outline=line, width=4)
    # battery indication
    if electric:
        d.rounded_rectangle((735,360,875,440), radius=8, outline=(45,106,79), width=5)
        d.text((755,378), "80 V", fill=(45,106,79), font=font(34, True))
    # dimension arrows
    d.line((175,575,1125,575), fill=(90,100,110), width=3)
    d.polygon([(175,575),(195,565),(195,585)], fill=(90,100,110))
    d.polygon([(1125,575),(1105,565),(1105,585)], fill=(90,100,110))
    d.text((560,585), "overall envelope - confirm by mast/configuration", fill=(90,100,110), font=font(24))
    d.text((75,35), title, fill=line, font=font(38, True))
    d.text((75,85), "Conceptual side elevation - not to scale", fill=(90,100,110), font=font(24))
    im.save(path, quality=95)

def draw_sideload(path):
    im = Image.new("RGB", (1500, 650), "white")
    d = ImageDraw.Draw(im); line=(23,54,93); fill=(232,238,245)
    d.line((80,520,1420,520), fill=(165,175,185), width=3)
    d.rounded_rectangle((230,300,1250,485), radius=25, outline=line, width=7, fill=fill)
    d.rectangle((700,155,950,300), outline=line, width=7, fill=(245,247,250))
    d.line((720,155,720,485), fill=line, width=7); d.line((930,155,930,485), fill=line, width=7)
    d.rectangle((470,335,650,450), outline=(45,106,79), width=5)
    d.text((493,372), "BATTERY", fill=(45,106,79), font=font(27, True))
    for x in [350,1120]:
        d.ellipse((x-65,440-65,x+65,440+65), fill=(40,45,50), outline=line, width=4)
        d.ellipse((x-30,410,x+30,470), fill="white", outline=line, width=4)
    d.line((810,325,810,95), fill=line, width=9); d.line((860,325,860,95), fill=line, width=9)
    d.line((810,255,590,255), fill=line, width=9)
    d.text((75,35), "Combilift C12000ET - multidirectional electric sideloader", fill=line, font=font(36, True))
    d.text((75,85), "Conceptual side elevation - not to scale", fill=(90,100,110), font=font(24))
    d.line((230,575,1250,575), fill=(90,100,110), width=3)
    d.text((540,585), "platform / chassis envelope", fill=(90,100,110), font=font(24))
    im.save(path, quality=95)

def draw_lowboy(path):
    im = Image.new("RGB", (1500, 650), "white")
    d=ImageDraw.Draw(im); line=(23,54,93); fill=(232,238,245)
    d.line((60,520,1440,520), fill=(165,175,185), width=3)
    d.polygon([(100,300),(330,300),(390,390),(1120,390),(1200,290),(1390,290),(1390,430),(1180,430),(1110,470),(380,470),(300,430),(100,430)], outline=line, fill=fill)
    for x in [1110,1210,1310]:
        d.ellipse((x-55,445-55,x+55,445+55), fill=(40,45,50), outline=line, width=4)
        d.ellipse((x-25,420,x+25,470), fill="white", outline=line, width=3)
    d.text((75,35), "Trail King 3-axle lowboy - proposed backhoe transport substitute", fill=line, font=font(34, True))
    d.text((75,85), "Conceptual side elevation - not to scale", fill=(90,100,110), font=font(24))
    d.line((390,560,1120,560), fill=(90,100,110), width=3)
    d.text((650,570), "main deck", fill=(90,100,110), font=font(24))
    im.save(path, quality=95)

draw_forklift(ASSET/"clark_s30.png", "CLARK S30 DS - diesel pneumatic forklift")
draw_forklift(ASSET/"taylor_gt60.png", "Taylor GT-60 - heavy-duty pneumatic forklift")
draw_forklift(ASSET/"clark_gex30.png", "CLARK GEX30 - 80 V electric forklift", electric=True)
draw_forklift(ASSET/"taylor_gt155d.png", "Taylor GT-155D - diesel heavy-duty forklift")
draw_sideload(ASSET/"combilift.png")
draw_lowboy(ASSET/"trailking.png")

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(8); normal.paragraph_format.line_spacing = 1.333
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for name, size, before, after, color in [("Title",28,0,8,NAVY),("Subtitle",14,0,16,GRAY),("Heading 1",16,18,10,BLUE),("Heading 2",13,12,6,BLUE),("Heading 3",12,8,4,NAVY)]:
    st=styles[name]; st.font.name="Calibri"; st._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); st._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri")
    st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=name.startswith("Heading") or name=="Title"
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

def shade(cell, color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"),color)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tc=cell._tc.get_or_add_tcPr(); tcMar=tc.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar=OxmlElement("w:tcMar"); tc.append(tcMar)
    for tag,val in [("top",top),("start",start),("bottom",bottom),("end",end)]:
        n=tcMar.find(qn(f"w:{tag}"))
        if n is None: n=OxmlElement(f"w:{tag}"); tcMar.append(n)
        n.set(qn("w:w"),str(val)); n.set(qn("w:type"),"dxa")

def set_table_geometry(table, widths):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT
    tblPr=table._tbl.tblPr
    tblW=tblPr.find(qn("w:tblW"))
    if tblW is None: tblW=OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"),str(sum(widths))); tblW.set(qn("w:type"),"dxa")
    ind=tblPr.find(qn("w:tblInd"))
    if ind is None: ind=OxmlElement("w:tblInd"); tblPr.append(ind)
    ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa")
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        gc=OxmlElement("w:gridCol"); gc.set(qn("w:w"),str(w)); grid.append(gc)
    for row in table.rows:
        for cell,w in zip(row.cells,widths):
            cell.width=Inches(w/1440); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)
            tcW=cell._tc.get_or_add_tcPr().find(qn("w:tcW")); tcW.set(qn("w:w"),str(w)); tcW.set(qn("w:type"),"dxa")

def table(rows, widths, header=True, font_size=9):
    t=doc.add_table(rows=len(rows), cols=len(widths)); t.style="Table Grid"; set_table_geometry(t,widths)
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            c=t.cell(i,j); c.text=str(val)
            if i==0 and header: shade(c,NAVY)
            for p in c.paragraphs:
                p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.05; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name="Calibri"; r.font.size=Pt(font_size); r.font.color.rgb=RGBColor.from_string(WHITE if i==0 and header else INK); r.bold=(i==0 and header)
    if header:
        trPr=t.rows[0]._tr.get_or_add_trPr(); hdr=OxmlElement("w:tblHeader"); hdr.set(qn("w:val"),"true"); trPr.append(hdr)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def add_para(text="", bold=False, italic=False, color=INK, size=11, align=None, after=8):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    if align is not None: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.name="Calibri"; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    return p

def add_bullets(items):
    for item in items:
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.375); p.paragraph_format.first_line_indent=Inches(-.194); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.208
        p.add_run(item)

def callout(label, text, color=PALE):
    t=doc.add_table(rows=1,cols=1); set_table_geometry(t,[9360]); shade(t.cell(0,0),color); margins(t.cell(0,0),140,180,140,180)
    p=t.cell(0,0).paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(label+"  "); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

def page_break(): doc.add_page_break()

def add_picture(path, width=6.45):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(4)
    r=p.add_run(); r.add_picture(str(path), width=Inches(width))

def add_caption(text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10)
    r=p.add_run(text); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(GRAY)

# Running header/footer
hp=sec.header.paragraphs[0]; hp.text="EQUIPMENT SUBSTITUTION PROPOSAL  |  DZSP21"; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs: r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(GRAY)
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
run=fp.add_run("Prepared for DZSP21  |  Commercial-in-Confidence  |  ")
fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); run._r.addnext(fld)
for r in fp.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)

# Cover - proposal_centerpiece pattern
add_para("[YOUR COMPANY NAME]", bold=True, color=GRAY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
p=doc.add_paragraph(style="Title"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Equipment Substitution Proposal")
p=doc.add_paragraph(style="Subtitle"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Forklifts and Lowboy Trailer")
add_para("Prepared exclusively for DZSP21", bold=True, color=NAVY, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
table([
    ["Prepared for", "DZSP21"],
    ["Prepared by", "[Your Company Name]"],
    ["Proposal date", "August 20, 2026"],
    ["Proposal reference", "[Insert reference / solicitation number]"],
    ["Validity", "Subject to final configuration, availability, and written quotation"],
], [2300,7060], header=False, font_size=10)
callout("PROPOSAL INTENT", "Provide equal-or-greater-capacity substitutions for the requested material-handling and backhoe-transport equipment while preserving the intended power type and operational role wherever practical.")
add_para("Commercial pricing, delivery schedule, warranty period, and acceptance terms are intentionally left as fillable placeholders pending DZSP21's final configuration and site requirements.", italic=True, color=GRAY, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

page_break()
add_heading("1. Executive proposal",1)
add_para("[Your Company Name] respectfully submits the following equipment substitution package to DZSP21. The proposed fleet uses CLARK, Taylor, Combilift, and Trail King equipment to satisfy or exceed the requested nominal capacity classes. The substitution plan consolidates the forklift requirement into three principal platforms, reducing operator familiarization and parts complexity while retaining a dedicated electric unit and a heavy-capacity alternative.")
callout("RECOMMENDED AWARD BASIS", "Accept the listed substitutions subject to confirmation of mast height, lowered height, fork length, tire type, attachments, emissions jurisdiction, UL/EE/DS classification, charging infrastructure, trailer payload, and transported backhoe dimensions.", "FFF7E6")
add_heading("2. Substitution schedule",1)
matrix=[
    ["Item","Requested equipment","Qty.","Proposed replacement","Qty.","Capacity position"],
    ["1","4K DS forklift","1","CLARK S30 DS","1","Exceeds 4K class"],
    ["2","4K forklift","1","Taylor GT-60","1","Exceeds 4K class"],
    ["3","4K EE forklift","1","CLARK GEX30 electric; EE/UL classification to be confirmed","1","Exceeds 4K class"],
    ["4","6K forklift","8","CLARK S30 DS / Taylor GT-60","6 / 2","Meets nominal 6K class*"],
    ["5","6K DS forklift","3","CLARK S30 DS","3","Meets nominal 6K class*"],
    ["6","10K forklift","1","Option A: Taylor GT-155D; Option B: Combilift C12000ET electric sideloader","1","Both exceed 10K class**"],
    ["7","Semi-trailer, 2-axle, backhoe","1","Trail King 3-axle lowboy (TK102HDG basis or equivalent available unit)","1","Higher axle count / heavy-haul configuration"],
]
table(matrix,[620,1700,520,3390,610,2520],True,8.2)
add_para("* Capacity must be confirmed at the required 24-inch load center and specified lift height. ** C12000ET capacity and 53-inch platform depth are based on the requirement information supplied; factory model/configuration confirmation is required.", italic=True, color=GRAY, size=8.5)
add_heading("3. Quantity summary",2)
table([
    ["Proposed model","Total quantity","Primary assignment"],
    ["CLARK S30 DS","10","4K DS and 6K/6K DS requirements"],
    ["Taylor GT-60","3","4K and 6K requirements"],
    ["CLARK GEX30","1","4K electric requirement"],
    ["Taylor GT-155D or Combilift C12000ET","1","10K requirement - select one"],
    ["Trail King 3-axle lowboy","1","Backhoe transport"],
],[3000,1500,4860],True,9)

page_break()
add_heading("4. CLARK S30 DS - proposed for Items 1, 4, and 5",1)
add_picture(ASSET/"clark_s30.png"); add_caption("Figure 1. Conceptual CLARK S30 DS side elevation; final mast and fork geometry will vary.")
table([
    ["Specification","Proposed basis"],
    ["Truck type","Counterbalanced pneumatic-tire forklift; diesel configuration"],
    ["Nominal capacity","Approximately 6,000-6,600 lb class; final rating at 24-in load center to be confirmed"],
    ["Manufacturer family rating","S20-S35 family: 2,000-3,500 kg at 500 mm load center"],
    ["Power","Diesel; published family data references 46.0 kW diesel power"],
    ["Mast / lift height","Two- or three-stage mast selected to DZSP21 clearance and lift requirements"],
    ["Tires","Pneumatic or superelastic, subject to operating surface"],
    ["Safety / classification","DS/UL construction and required labels to be confirmed on offered serial-numbered unit"],
],[2700,6660],True,9)
add_heading("Substitution rationale",2)
add_bullets(["One platform covers the 4K DS and 6K/6K DS roles, simplifying fleet support.","Nominal capacity exceeds the 4K requirement and is aligned to the 6K class.","Wet-disc brake and operator-assist features are available within the S-Series family; final options will be listed on the quotation."])

page_break()
add_heading("5. Taylor GT-60 - proposed for Items 2 and 4",1)
add_picture(ASSET/"taylor_gt60.png"); add_caption("Figure 2. Conceptual Taylor GT-60 side elevation; not to scale.")
table([
    ["Specification","Published / proposed basis"],
    ["Rated capacity","6,000 lb at 24-in load center"],
    ["Wheelbase","66.9 in"],
    ["Engine","HMC Theta 2.4 L, 65 hp Tier 4 Final LPG standard; Isuzu 2.2 L, 62 hp Tier 4 Final diesel available"],
    ["Transmission / axle","Single-speed powershift with inching; planetary drive axle with wet-disc brakes"],
    ["Mast","Two- and three-stage uprights; published family lift heights up to 217 in"],
    ["Standard GT-60 mast basis","130-in maximum fork height two-stage mast; approximately 87.9-in lowered height"],
    ["Standard forks","1.75 x 5 x 42 in, pin-mounted; alternate length available subject to approval"],
    ["Operator station","Overhead guard, suspension seat, LCD dash, adjustable steering column"],
],[2700,6660],True,9)
add_heading("Substitution rationale",2)
add_bullets(["Rated for the 6K requirement at a 24-inch load center.","Provides substantial reserve over the requested 4K class.","Heavy-duty construction is suitable for outdoor and mixed-yard duty, subject to tire selection."])

page_break()
add_heading("6. CLARK GEX30 - proposed for Item 3",1)
add_picture(ASSET/"clark_gex30.png"); add_caption("Figure 3. Conceptual CLARK GEX30 electric forklift side elevation; charger not shown.")
table([
    ["Specification","Proposed basis"],
    ["Truck type","Four-wheel counterbalanced electric forklift"],
    ["Nominal capacity","3,000 kg / approximately 6,600 lb class; verify rating at required load center and lift height"],
    ["Electrical system","80 V traction system"],
    ["Mast / forks","Configured to DZSP21 lift-height, lowered-height, free-lift, and fork-length requirements"],
    ["Battery / charger","Battery chemistry, amp-hour capacity, connector, and charger input voltage to be confirmed"],
    ["EE requirement","EE/UL listing is not assumed from model name; written manufacturer certification and truck data plate must confirm compliance"],
],[2700,6660],True,9)
callout("CRITICAL COMPLIANCE HOLD POINT", "Do not accept the GEX30 as an EE substitute until the offered serial-numbered truck, battery, and charger package carries the classification required by DZSP21 and the intended hazardous-location policy.", "FCE8E6")

page_break()
add_heading("7. 10K forklift alternatives - Item 6",1)
add_heading("Option A - Taylor GT-155D",2)
add_picture(ASSET/"taylor_gt155d.png",6.2)
table([
    ["Specification","Published / proposed basis"],
    ["Rated capacity","15,500 lb at 24-in load center"],
    ["Wheelbase","88.6 in"],
    ["Power","Diesel configuration proposed; engine/emissions package to be confirmed for destination"],
    ["Application fit","Conventional front-loading heavy-duty forklift; significant capacity reserve over 10K"],
],[2700,6660],True,9)
add_heading("Option B - Combilift C12000ET electric sideloader",2)
add_picture(ASSET/"combilift.png",6.2)
table([
    ["Specification","Requirement / proposed basis"],
    ["Rated capacity","12,000 lb at 24-in load center - based on supplied requirement; factory confirmation required"],
    ["Platform depth","53 in - based on supplied requirement; confirm exact model drawing"],
    ["Power / handling","Electric, multidirectional sideloader; suited to long loads and narrow-aisle movement"],
    ["Battery / charger","To be sized for duty cycle and facility electrical service"],
],[2700,6660],True,9)
callout("SELECTION GUIDANCE", "Choose the Taylor GT-155D for conventional front-lift work and maximum capacity reserve. Choose the Combilift when long-load handling, side loading, indoor emissions, or aisle maneuverability is the controlling requirement.")

page_break()
add_heading("8. Trail King 3-axle lowboy - proposed for Item 7",1)
add_picture(ASSET/"trailking.png"); add_caption("Figure 4. Conceptual Trail King 3-axle detachable-gooseneck lowboy; final offered unit may vary.")
table([
    ["Specification","TK102HDG basis / proposed requirement"],
    ["Trailer type","Hydraulic detachable gooseneck lowboy, semi-trailer"],
    ["Axles","Three 25,000-lb axles; third axle air lift"],
    ["Rated capacity","102,000 lb in 12 ft (manufacturer model basis)"],
    ["Main deck length","25 ft 9 in"],
    ["Axle spacing","55 in"],
    ["Loaded fifth-wheel height","50 in"],
    ["Decking","1.75-in apitong raised decking listed for TK102HDG"],
    ["Final transport check","Backhoe operating weight, axle loads, width, height, wheelbase, ground clearance, tie-down points, tractor compatibility, and local permit limits"],
],[2700,6660],True,9)
add_heading("Substitution rationale",2)
add_bullets(["The three-axle configuration provides a heavy-haul platform in place of the requested two-axle backhoe trailer.","Detachable gooseneck supports front loading of construction equipment.","Final acceptance must be based on the actual backhoe and prime mover, not trailer capacity alone."])

page_break()
add_heading("9. Configuration, compliance, and acceptance",1)
add_para("The proposal defines suitable model families and capacity positions. It is not a final engineered configuration. DZSP21 and the supplier should complete the following schedule before order release.")
table([
    ["Acceptance item","DZSP21 requirement / supplier response"],
    ["Rated load","Confirm capacity at actual load center, lift height, attachment derate, and side-shift/fork-positioner configuration."],
    ["Mast envelope","State maximum fork height, free lift, lowered mast height, raised mast height, and overhead-guard height."],
    ["Forks / carriage","Confirm fork length, section, spread, carriage class, load backrest, and attachments."],
    ["Power / emissions","Confirm diesel/LPG/electric choice, emissions jurisdiction, fuel type, battery duty cycle, and charger input."],
    ["Safety classification","Provide written UL/EE/DS certification where required; confirm lighting, alarm, strobe, restraint, fire extinguisher, and data plate."],
    ["Environment","Confirm indoor/outdoor use, grades, floor loading, aisle width, weather, corrosion exposure, and tire compound."],
    ["Trailer compatibility","Provide backhoe and tractor data; complete axle-load, turning-clearance, deck, tiedown, and regulatory review."],
    ["Documentation","Operator manuals, service manuals, parts books, certificates, inspection records, and warranty terms."],
    ["Inspection / training","Pre-delivery inspection, site acceptance test, familiarization, and operator/maintenance training."],
],[3000,6360],True,8.8)
add_heading("10. Commercial schedule - to be completed",1)
table([
    ["Commercial item","Offer"],
    ["Total price","[Insert price and currency]"],
    ["Delivery","[Insert lead time / delivery location]"],
    ["Proposal validity","[Insert validity period]"],
    ["Warranty","[Insert coverage]"],
    ["Payment terms","[Insert terms]"],
    ["Freight / duties / taxes","[State inclusion or exclusion]"],
],[2600,6760],True,9)

page_break()
add_heading("11. Assumptions and qualifications",1)
add_bullets([
    "Model names in this proposal are interpreted as equipment families. The final serial-numbered units and option codes govern.",
    "No attachment derating, aisle study, floor-loading analysis, hazardous-location determination, or trailer route/permit analysis has been performed.",
    "The term DS is treated as a required safety/construction classification and must be validated by manufacturer documentation and the truck data plate.",
    "The CLARK GEX30 is proposed as an electric substitute; EE compliance remains an explicit hold point.",
    "The Combilift C12000ET capacity and 53-inch platform depth are retained from the supplied requirement because the exact factory technical sheet was not identified; written factory confirmation is required before acceptance.",
    "Availability may require an equivalent model-year configuration with equal or greater capacity, subject to DZSP21 approval.",
    "All drawings are conceptual, not to scale, and are not approved-for-construction or clearance drawings.",
])
add_heading("12. Reference sources",1)
sources=[
    ["Manufacturer","Reference"],
    ["CLARK","S20-S35 product page and published family technical data: https://clarkmheu.com/en/forklifts-with-diesel-or-lpg-drive/s20-25-30-35"],
    ["CLARK","GEX electric product-family information: https://clarkmheu.com/en/parts/cabins-rain-caps"],
    ["Taylor","GT Series brochure: https://taylorforklifts.com/assets/brochures/product/GT-Series.pdf"],
    ["Taylor","GT-60 / GT-66 brochure: https://taylorforklifts.com/assets/brochures/product/GT-60_GT-66_Brochure.pdf"],
    ["Trail King","TKHDG hydraulic detachable gooseneck specifications: https://www.trailking.com/products/tkhdg-hydraulic-detachable-gooseneck/"],
    ["Combilift","C-Series reference family sheet: https://combilift.com/wp-content/uploads/2020/10/C-Series-12000-Technical-UK-2018.pdf"],
]
table(sources,[1700,7660],True,8.1)
add_para("Source data is summarized for proposal development only. Manufacturer literature and specifications can change; the final quotation, certified capacity plate, and approved general-arrangement drawing control.", italic=True, color=GRAY, size=8.5)
add_heading("13. Acceptance / authorization",1)
table([
    ["For DZSP21","For [Your Company Name]"],
    ["Name: ______________________________\nTitle: _______________________________\nDate: ________________________________","Name: ______________________________\nTitle: _______________________________\nDate: ________________________________"],
],[4680,4680],True,9)

# Core properties
doc.core_properties.title="Equipment Substitution Proposal - DZSP21"
doc.core_properties.subject="Forklift and lowboy trailer substitutions"
doc.core_properties.author="[Your Company Name]"
doc.core_properties.keywords="DZSP21, forklift, substitution, CLARK, Taylor, Combilift, Trail King"
doc.save(OUT)
print(OUT)
