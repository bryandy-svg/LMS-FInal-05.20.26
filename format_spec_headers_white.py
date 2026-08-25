from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "LMS_0200-6999_DZSP21_Proposal_Three_Trailer_Options_Photos.docx"
OUTPUT = ROOT / "LMS_0200-6999_DZSP21_Proposal_Formatted_Trailer_Options.docx"
HEADER_FILL = "1F4E78"


def set_header_cell(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), HEADER_FILL)
    tc_pr.append(shading)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        tr_pr.append(repeat)


def main():
    doc = Document(SOURCE)
    formatted = 0
    for table in doc.tables:
        if not table.rows or table.rows[0].cells[0].text != "Specification":
            continue
        set_repeat_header(table.rows[0])
        for cell in table.rows[0].cells:
            set_header_cell(cell)
        formatted += 1

    if formatted != 8:
        raise RuntimeError(f"Expected 8 specification tables, formatted {formatted}")
    doc.save(OUTPUT)

    check = Document(OUTPUT)
    checked = 0
    failures = []
    for table in check.tables:
        if not table.rows or table.rows[0].cells[0].text != "Specification":
            continue
        checked += 1
        for cell in table.rows[0].cells:
            fill = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
            if fill is None or fill.get(qn("w:fill")) != HEADER_FILL:
                failures.append(f"table {checked}: header fill")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.color.rgb != RGBColor(255, 255, 255):
                        failures.append(f"table {checked}: font color")

    drawings = len(check.element.body.xpath('.//w:drawing'))
    print(f"output={OUTPUT}")
    print(f"spec_tables_formatted={checked}")
    print(f"format_failures={failures}")
    print(f"body_drawings={drawings}")
    if checked != 8 or failures or drawings != 8:
        raise RuntimeError("Post-save verification failed")


if __name__ == "__main__":
    main()
